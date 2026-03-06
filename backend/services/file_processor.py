"""
File Processing Service
Extracts text content from supported file formats for PII detection.

Supported formats: CSV, JSON, SQL, PDF, DOCX/DOC

Pipeline step: File Type Detection → File Parsing → Text/Data Extraction
"""
import json
import pandas as pd
import fitz  # PyMuPDF
from docx import Document
import sqlparse
from typing import List, Dict, Any


def extract_csv(filepath: str) -> List[Dict[str, Any]]:
    """Extract data from a CSV file as list of row dicts."""
    df = pd.read_csv(filepath, dtype=str, keep_default_na=False)
    return df.to_dict(orient="records")


def extract_json(filepath: str) -> Any:
    """Load and return JSON content."""
    with open(filepath, "r", encoding="utf-8") as f:
        return json.load(f)


def extract_pdf(filepath: str) -> str:
    """Extract all text from a PDF file."""
    doc = fitz.open(filepath)
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()
    return "\n".join(text_parts)


def extract_docx(filepath: str) -> str:
    """Extract text from paragraphs and tables in a DOCX file."""
    doc = Document(filepath)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def extract_sql(filepath: str) -> List[str]:
    """Extract values from INSERT/UPDATE statements in SQL dumps."""
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()

    statements = sqlparse.parse(content)
    values = []
    for stmt in statements:
        sql_str = str(stmt).strip()
        upper = sql_str.upper()
        if upper.startswith("INSERT"):
            idx = upper.find("VALUES")
            if idx != -1:
                vals_section = sql_str[idx + 6:].strip().rstrip(";")
                depth = 0
                current: list = []
                for ch in vals_section:
                    if ch == "(":
                        depth += 1
                        if depth == 1:
                            current = []
                            continue
                    elif ch == ")":
                        depth -= 1
                        if depth == 0:
                            row_str = "".join(current)
                            for val in row_str.split(","):
                                cleaned = val.strip().strip("'\"")
                                if cleaned:
                                    values.append(cleaned)
                            continue
                    if depth >= 1:
                        current.append(ch)
        elif upper.startswith("UPDATE"):
            set_idx = upper.find("SET")
            where_idx = upper.find("WHERE")
            if set_idx != -1:
                end = where_idx if where_idx != -1 else len(sql_str)
                set_clause = sql_str[set_idx + 3:end]
                for assignment in set_clause.split(","):
                    if "=" in assignment:
                        val = assignment.split("=", 1)[1].strip().strip("'\"")
                        if val:
                            values.append(val)
    return values


def detect_file_type(filename: str) -> str:
    """Detect file type from extension."""
    import os
    _, ext = os.path.splitext(filename)
    return ext.lower().lstrip(".")


def extract_text_from_file(filepath: str, file_type: str) -> str:
    """
    Master extraction function.
    Normalizes content from any supported format into a single text string
    suitable for PII detection.
    """
    file_type = file_type.lower().lstrip(".")

    if file_type == "csv":
        rows = extract_csv(filepath)
        lines = []
        for row in rows:
            line = " | ".join(f"{k}: {v}" for k, v in row.items())
            lines.append(line)
        return "\n".join(lines)

    elif file_type == "json":
        data = extract_json(filepath)
        return json.dumps(data, indent=2, ensure_ascii=False)

    elif file_type == "pdf":
        return extract_pdf(filepath)

    elif file_type in ("docx", "doc"):
        return extract_docx(filepath)

    elif file_type == "sql":
        values = extract_sql(filepath)
        return "\n".join(values)

    else:
        raise ValueError(f"Unsupported file type: {file_type}")
