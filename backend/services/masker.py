"""
Masking Engine
Replaces detected PII with tokens and reconstructs sanitized output files
in the original format (CSV, JSON, SQL, PDF, DOCX).

Pipeline step: Tokenized PII mapping → Sanitized File Generation
"""
import json
import os
import pandas as pd
from typing import Dict
from config import settings


def mask_text(text: str, mapping: Dict[str, str]) -> str:
    """
    Replace all PII occurrences in text with their token replacements.
    Processes longest matches first to avoid partial replacements.
    """
    sorted_originals = sorted(mapping.keys(), key=len, reverse=True)
    result = text
    for original in sorted_originals:
        result = result.replace(original, mapping[original])
    return result


def mask_csv(filepath: str, mapping: Dict[str, str], output_path: str):
    """Mask PII in a CSV file and write sanitized version."""
    df = pd.read_csv(filepath, dtype=str, keep_default_na=False)
    for col in df.columns:
        df[col] = df[col].apply(lambda val: mask_text(str(val), mapping))
    df.to_csv(output_path, index=False)


def mask_json(filepath: str, mapping: Dict[str, str], output_path: str):
    """Mask PII in a JSON file, preserving nested structure."""
    with open(filepath, "r", encoding="utf-8") as f:
        data = json.load(f)

    def _mask_recursive(obj):
        if isinstance(obj, str):
            return mask_text(obj, mapping)
        elif isinstance(obj, dict):
            return {k: _mask_recursive(v) for k, v in obj.items()}
        elif isinstance(obj, list):
            return [_mask_recursive(item) for item in obj]
        return obj

    masked_data = _mask_recursive(data)
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(masked_data, f, indent=2, ensure_ascii=False)


def mask_sql(filepath: str, mapping: Dict[str, str], output_path: str):
    """Mask PII in a SQL dump file."""
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()
    masked = mask_text(content, mapping)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(masked)


def mask_docx(filepath: str, mapping: Dict[str, str], output_path: str):
    """Mask PII in a DOCX file, preserving formatting."""
    from docx import Document
    doc = Document(filepath)

    for para in doc.paragraphs:
        for run in para.runs:
            run.text = mask_text(run.text, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = mask_text(run.text, mapping)

    doc.save(output_path)


def mask_pdf(filepath: str, mapping: Dict[str, str], output_path: str):
    """
    Mask PII in a PDF. Extracts text, applies masking, and creates
    a new text-based PDF (original layout may not be preserved).
    """
    import fitz
    doc = fitz.open(filepath)
    new_doc = fitz.open()

    for page in doc:
        text = page.get_text()
        masked = mask_text(text, mapping)

        new_page = new_doc.new_page(
            width=page.rect.width, height=page.rect.height,
        )
        new_page.insert_textbox(
            fitz.Rect(50, 50, page.rect.width - 50, page.rect.height - 50),
            masked, fontsize=10, fontname="helv",
        )

    new_doc.save(output_path)
    new_doc.close()
    doc.close()


def mask_file(filepath: str, file_type: str, mapping: Dict[str, str], output_path: str):
    """Route to the appropriate masking function based on file type."""
    file_type = file_type.lower().lstrip(".")

    if file_type == "csv":
        mask_csv(filepath, mapping, output_path)
    elif file_type == "json":
        mask_json(filepath, mapping, output_path)
    elif file_type in ("docx", "doc"):
        mask_docx(filepath, mapping, output_path)
    elif file_type == "pdf":
        mask_pdf(filepath, mapping, output_path)
    elif file_type == "sql":
        mask_sql(filepath, mapping, output_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
